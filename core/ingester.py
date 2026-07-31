"""
数据摄入模块
负责解析 PDF/Word/Excel/CSV 固井资料，输出标准化 Document 对象
"""

from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field

from loguru import logger


@dataclass
class CementDocument:
    """固井文档标准化结构"""
    content: str                          # 文本内容
    source: str                           # 来源文件名
    doc_type: str = "unknown"             # 文档类型：report / data / log
    metadata: dict = field(default_factory=dict)  # 元数据（井名、日期、井段等）


class DocumentIngester:
    """固井文档摄入器

    支持格式：
    - PDF（固井施工报告、完井报告）—— 重点支持
    - Word（固井施工报告、技术文档）
    - Excel（固井数据表、施工参数记录）
    - CSV（结构化数据）
    """

    def __init__(self, raw_dir: Path):
        self.raw_dir = Path(raw_dir)
        self.raw_dir.mkdir(parents=True, exist_ok=True)

    def ingest_all(self) -> List[CementDocument]:
        """扫描 raw_dir 下所有支持的文件并解析"""
        documents = []
        for file_path in self.raw_dir.rglob("*"):
            if file_path.is_dir():
                continue
            suffix = file_path.suffix.lower()
            if suffix == ".pdf":
                docs = self._parse_pdf(file_path)
            elif suffix in (".docx", ".doc"):
                docs = self._parse_word(file_path)
            elif suffix in (".xlsx", ".xls"):
                docs = self._parse_excel(file_path)
            elif suffix == ".csv":
                docs = self._parse_csv(file_path)
            else:
                logger.debug(f"跳过不支持的文件格式: {file_path}")
                continue
            documents.extend(docs)
            logger.info(f"解析完成: {file_path.name} → {len(docs)} 个文档片段")

        logger.info(f"共解析 {len(documents)} 个文档片段")
        return documents

    def ingest_file(self, file_path: Path) -> List[CementDocument]:
        """解析单个文件"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self._parse_word(file_path)
        elif suffix in (".xlsx", ".xls"):
            return self._parse_excel(file_path)
        elif suffix == ".csv":
            return self._parse_csv(file_path)
        else:
            logger.warning(f"不支持的文件格式: {suffix}")
            return []

    def _parse_pdf(self, file_path: Path) -> List[CementDocument]:
        """解析 PDF 文件（固井施工报告）

        使用 PyMuPDF 提取文本；当页面文本过短（疑似扫描件）时，
        自动调用 PaddleOCR VL 进行图像识别兜底。
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("请安装 PyMuPDF: pip install pymupdf")
            return []

        # 延迟导入 OCR 引擎配置
        from config import settings as _settings
        ocr_enabled = _settings.ocr.enabled
        min_text_len = _settings.ocr.min_text_length

        documents = []
        ocr_used_pages = 0

        try:
            doc = fitz.open(str(file_path))
            full_text = ""

            for page_idx, page in enumerate(doc):
                page_text = page.get_text()

                # 文本过短 → 尝试 OCR 兜底（扫描件）
                if len(page_text.strip()) < min_text_len and ocr_enabled:
                    ocr_text = self._ocr_page(page)
                    if ocr_text:
                        page_text = ocr_text
                        ocr_used_pages += 1
                        logger.debug(f"第 {page_idx + 1} 页使用 OCR 识别，文本长度: {len(ocr_text)}")

                full_text += page_text

            # 提取元数据（从文本中匹配常见固井字段）
            metadata = self._extract_metadata(full_text, file_path)
            metadata["file_path"] = str(file_path)
            metadata["total_pages"] = len(doc)
            if ocr_used_pages > 0:
                metadata["ocr_used"] = True
                metadata["ocr_pages"] = ocr_used_pages

            # 按段落拆分（空行分隔）
            paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]

            if paragraphs:
                # 整体作为一个文档（保留上下文完整性）
                documents.append(CementDocument(
                    content=full_text,
                    source=file_path.name,
                    doc_type="report",
                    metadata=metadata,
                ))

            doc.close()
            if ocr_used_pages > 0:
                logger.info(f"PDF OCR 完成: {file_path.name}，{ocr_used_pages}/{len(doc)} 页使用 OCR")
        except Exception as e:
            logger.error(f"PDF 解析失败 {file_path}: {e}")

        return documents

    def _ocr_page(self, page) -> str:
        """对单个 PDF 页面执行 OCR 识别

        将页面渲染为图片后交给 PaddleOCR VL 识别
        """
        try:
            from PIL import Image
            import io
            from pipeline.ocr_engine import ocr_page
        except ImportError as e:
            logger.warning(f"OCR 依赖缺失: {e}")
            return ""

        try:
            # 渲染页面为图片（200 DPI，平衡清晰度和速度）
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            return ocr_page(img)
        except Exception as e:
            logger.warning(f"页面 OCR 失败: {e}")
            return ""

    def _parse_word(self, file_path: Path) -> List[CementDocument]:
        """解析 Word 文件（.docx / .doc）

        使用 python-docx 解析 .docx；
        .doc 旧格式尝试通过 pywin32（Windows）转换后解析
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("请安装 python-docx: pip install python-docx")
            return []

        documents = []
        try:
            suffix = file_path.suffix.lower()
            actual_path = file_path

            # .doc 旧格式需要先转换
            if suffix == ".doc":
                converted = self._convert_doc_to_docx(file_path)
                if converted is None:
                    logger.error(f".doc 转换失败，跳过: {file_path}")
                    return []
                actual_path = converted

            doc = DocxDocument(str(actual_path))

            # 提取所有段落文本
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            # 提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_text.append("\t".join(cells))
                if table_text:
                    full_text += "\n" + "\n".join(table_text)

            if not full_text.strip():
                logger.warning(f"Word 文件内容为空: {file_path}")
                return []

            metadata = self._extract_metadata(full_text, file_path)
            metadata["file_path"] = str(file_path)
            metadata["paragraph_count"] = len(paragraphs)
            metadata["table_count"] = len(doc.tables)

            documents.append(CementDocument(
                content=full_text,
                source=file_path.name,
                doc_type="report",
                metadata=metadata,
            ))

            # 清理临时转换文件
            if suffix == ".doc" and actual_path != file_path:
                actual_path.unlink(missing_ok=True)

        except Exception as e:
            logger.error(f"Word 解析失败 {file_path}: {e}")

        return documents

    def _convert_doc_to_docx(self, doc_path: Path) -> Optional[Path]:
        """通过 pywin32 将 .doc 转换为 .docx（仅 Windows）"""
        try:
            import win32com.client
            import pythoncom
        except ImportError:
            logger.error("解析 .doc 需要 pywin32: pip install pywin32")
            return None

        docx_path = doc_path.with_suffix(".docx")
        word = None
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(doc_path.resolve()))
            doc.SaveAs2(str(docx_path.resolve()), FileFormat=16)  # 16 = docx
            doc.Close()
            return docx_path
        except Exception as e:
            logger.error(f".doc 转换失败: {e}")
            return None
        finally:
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def _parse_excel(self, file_path: Path) -> List[CementDocument]:
        """解析 Excel 文件（固井数据表）"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("请安装 pandas + openpyxl")
            return []

        documents = []
        try:
            xls = pd.ExcelFile(str(file_path))
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                if df.empty:
                    continue

                # 转为文本描述
                content = f"【工作表: {sheet_name}】\n"
                content += df.to_string(index=False)

                metadata = self._extract_metadata(content, file_path)
                metadata["sheet_name"] = sheet_name
                metadata["row_count"] = len(df)
                metadata["columns"] = list(df.columns)

                documents.append(CementDocument(
                    content=content,
                    source=file_path.name,
                    doc_type="data",
                    metadata=metadata,
                ))
        except Exception as e:
            logger.error(f"Excel 解析失败 {file_path}: {e}")

        return documents

    def _parse_csv(self, file_path: Path) -> List[CementDocument]:
        """解析 CSV 文件"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("请安装 pandas")
            return []

        documents = []
        try:
            df = pd.read_csv(str(file_path))
            content = df.to_string(index=False)

            metadata = self._extract_metadata(content, file_path)
            metadata["row_count"] = len(df)
            metadata["columns"] = list(df.columns)

            documents.append(CementDocument(
                content=content,
                source=file_path.name,
                doc_type="data",
                metadata=metadata,
            ))
        except Exception as e:
            logger.error(f"CSV 解析失败 {file_path}: {e}")

        return documents

    def _extract_metadata(self, text: str, file_path: Path) -> dict:
        """从文本中提取固井相关元数据（井名、日期、井段等）

        使用简单的正则匹配，可根据实际数据格式扩展
        """
        import re

        metadata = {"filename": file_path.name}

        # 匹配井名（常见格式：XX-1、XX1-H1、XX-1-2）
        well_pattern = re.search(r'井\s*名[：:]\s*(\S+)', text)
        if well_pattern:
            metadata["well_name"] = well_pattern.group(1)
        else:
            # 尝试从文件名提取
            metadata["well_name"] = file_path.stem

        # 匹配日期
        date_pattern = re.search(r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})', text)
        if date_pattern:
            metadata["date"] = f"{date_pattern.group(1)}-{date_pattern.group(2).zfill(2)}-{date_pattern.group(3).zfill(2)}"

        # 匹配井深
        depth_pattern = re.search(r'(?:井深|完钻井深)[：:]\s*([\d.]+)\s*[mM米]', text)
        if depth_pattern:
            metadata["well_depth"] = float(depth_pattern.group(1))

        # 匹配井段
        section_pattern = re.search(r'(?:固井井段|封固井段)[：:]\s*([\d.]+)\s*[-~—]\s*([\d.]+)', text)
        if section_pattern:
            metadata["cement_section"] = f"{section_pattern.group(1)}-{section_pattern.group(2)}m"

        return metadata
